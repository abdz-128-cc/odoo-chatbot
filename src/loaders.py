from __future__ import annotations
import os, fitz
from typing import List
from langchain_core.documents import Document
from docx import Document as DocxDocument

def load_pdf_with_pages(path: str) -> List[Document]:
    """
    Loads a PDF file and creates documents for each page with metadata.

    Args:
        path: The path to the PDF file.

    Returns:
        A list of documents, one per page with content.
    """
    doc = fitz.open(path)
    documents = []
    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():  # Only add pages with actual content
            meta = {
                "source": "handbook",
                "path": path,
                "page": page_num + 1  # Add page number to metadata
            }
            documents.append(Document(page_content=text, metadata=meta))
    return documents

def load_docx(path: str) -> str:
    """
    Loads text content from a DOCX file.

    Args:
        path: The path to the DOCX file.

    Returns:
        The concatenated text from all paragraphs.
    """
    d = DocxDocument(path)
    return "\n".join([p.text for p in d.paragraphs])

def walk_docs(root: str) -> List[Document]:
    """
    Walks a directory and loads supported documents (PDF, DOCX).

    Args:
        root: The root directory to walk.

    Returns:
        A list of loaded documents.
    """
    docs = []
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            p = os.path.join(dirpath, fn)
            ext = os.path.splitext(fn)[1].lower()
            if ext == ".pdf":
                docs.extend(load_pdf_with_pages(p))
            elif ext in (".docx",):
                text = load_docx(p)
                if text.strip():
                    meta = {"source": "handbook", "path": p}
                    docs.append(Document(page_content=text, metadata=meta))
            else:
                continue
    return docs
